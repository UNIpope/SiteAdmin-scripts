from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import pandas as pd
import numpy as np
import os 

def fillindocx(fname, oname, filld):
    document = Document(fname)

    # make head style 
    styles = document.styles
    styles.add_style('head', WD_STYLE_TYPE.PARAGRAPH)

    # set style 
    stylefoot = document.styles['head']
    fontfoot = stylefoot.font
    fontfoot.name = 'Arial'
    fontfoot.size = Pt(12)

    # this is the worst thing ive ever seen 
    for tb in document.tables:
        for r in tb.rows:
            for c in r.cells:
                for para in c.paragraphs:
                    for i in filld:
                        if i in para.text:
                            txt = para.text
                            otxt = txt.replace(i, filld[i])
                            para.text = otxt

    for para in document.paragraphs:
        for i in filld:
            if i in para.text:
                txt = para.text
                otxt = txt.replace(i, filld[i])
                para.text = otxt
                para.style = document.styles['head']

                if i == filld["12NAME"]:
                    fontfoot.bold = True
                    

    document.save(oname)

csvin = r'siteadmin\wordtemplating\Frank\nameadde.csv'
df = pd.read_csv(csvin, engine='python')
df = df.apply(lambda x: x.str.strip())
df = df.apply(lambda x: x.str.upper())

# file paths
opath = r"C:\Users\jduggan\Desktop\script\siteadmin\wordtemplating\Frank\out"
tpath = r'siteadmin\wordtemplating\Frank\Template - Frankfurt Movement.docx'


for _, row in df.iterrows():
    adde = row["Address"].split(", ")[:-1]

    adtop = ",\n".join(adde)
    admain = ",".join(adde)
    name = row["Name"]

    rep = {"12NAME":name, "ADDE3": adtop, "3ADDE":admain}
    path = opath + "\\" + name + " - Frankfurt Movement.docx"

    print(rep)

    fillindocx(tpath, path, rep)
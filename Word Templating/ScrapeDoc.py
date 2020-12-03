from docx import Document
import pandas as pd

fname = r'C:\Users\jduggan\Desktop\script\siteadmin\wordtemplating\PAR\in\Amandine Joubert - Movement Letter.docx'
tname = r'C:\Users\jduggan\Desktop\script\siteadmin\wordtemplating\PAR\template.docx'

document = Document(fname)
template = Document(tname)

flookup = {"Fonctions" : "ROLE", "Nom" : "1NAME", "Prénom" : "2SURNAME", "Date de naissance" : "DOBIRTH",
            "Lieu de naissance" : "POBIRTH", "Adresse du domicile" : "PORES", 
            "Lieux d’exercice de l’activité professionnelle": "POWORK"}

odic = {}
for para in document.paragraphs:
    if ":" in para.text:
        k = para.text.split(":")[0].strip()
        o = para.text.split(":")[1].strip()

        if k in flookup.keys():
            print(flookup[k])
            odic[flookup[k]] = o

print(odic)


def scrape(fname):
    document = Document(fname)
    flookup = {"Fonctions" : "ROLE", "Nom" : "1NAME", "Prénom" : "2SURNAME", "Date de naissance" : "DOBIRTH",
                "Lieu de naissance" : "POBIRTH", "Adresse du domicile" : "PORES", 
                "Lieux d’exercice de l’activité professionnelle": "POWORK"}

    odic = {}
    for para in document.paragraphs:
        if ":" in para.text:
            k = para.text.split(":")[0].strip()
            o = para.text.split(":")[1].strip()

            if k in flookup.keys():
                print(flookup[k])
                odic[flookup[k]] = o
    
    odic["Fname"] = fname.split('\\')[-1]
    return odic

fname = r'C:\Users\jduggan\Desktop\script\siteadmin\wordtemplating\PAR\in\Amandine Joubert - Movement Letter.docx'
print(scrape(fname))